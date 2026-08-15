"""Build generic biomedical EN/ZH manuscript sources and DOCX deliverables.

The source results remain unchanged. This script converts the audited English
draft and its Chinese parallel manuscript into journal-neutral, typeset files
with a title page, declarations, references, main tables, and main figures.
"""

from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


EN_TITLE = (
    "Cohort-Aware Descriptive Analysis of Directionality and Heterogeneity of "
    "Nucleus Pulposus Transcriptional Programs across Public Human "
    "Intervertebral Disc Degeneration Datasets"
)
EN_SHORT_TITLE = "Cohort-aware IVDD program analysis"
ZH_TITLE = "公开人类椎间盘退变队列中髓核转录模块方向性与异质性的队列感知描述性分析"
ZH_SHORT_TITLE = "IVDD 髓核转录模块的队列级异质性"

EN_ABSTRACT = """Public single-cell intervertebral disc degeneration (IVDD) datasets often contain many cells but few independently observed donors, and clinical and processing differences complicate cross-cohort comparison. We conducted a cohort-aware descriptive analysis of four nucleus pulposus (NP) transcriptional programs: extracellular matrix (ECM)/collagen remodeling, inflammatory/nuclear factor-kappa B (NF-κB) response, hypoxia/oxidative stress, and disc-matrix homeostasis. The observation unit was the donor or, where needed, an explicitly labelled presumed donor, sample, or library key; cells were nested observations. Four audited NP cohorts contributed to the frozen default summary, which contained 20 effects: 16 NP effects and 4 exploratory annulus fibrosus (AF) effects from the GSE230809 parent project. We estimated cohort-specific differences between groups recorded as higher versus lower severity with descriptive Welch and donor or library bootstrap 95% intervals and leave-one-key-out stability. Exactly 55 of 55 score-to-ledger identity matches passed. Hypoxia/oxidative stress was the only NP program with positive point estimates in all four default cohorts (0.1776, 0.2381, 0.0882, and 0.4346), although every Welch interval included zero. The other three programs showed discordant directions. Separate exploratory standardized syntheses did not alter the default analysis. In a post hoc six-cohort expansion, the hypoxia standardized mean difference was 0.7694 (95% CI 0.1706 to 1.3682); after source-family replacement, it was 0.5746 (95% CI -0.7231 to 1.8723). These results suggest a descriptive hypoxia-related direction for future study amid substantial cohort-level heterogeneity. They do not establish a universal IVDD program, biological mechanism, biomarker, or therapeutic target."""

EN_KEYWORDS = (
    "intervertebral disc degeneration; nucleus pulposus; public transcriptomic "
    "data; single-cell transcriptomics; cohort heterogeneity; reproducibility"
)

MODULE_LABELS_EN = {
    "ECM / collagen remodeling": "ECM/collagen remodeling",
    "Inflammatory / NF-kB": "Inflammatory/NF-κB",
    "Hypoxia / oxidative stress": "Hypoxia/oxidative stress",
    "Disc matrix homeostasis": "Disc-matrix homeostasis",
}
MODULE_LABELS_ZH = {
    "ECM / collagen remodeling": "细胞外基质/胶原重塑",
    "Inflammatory / NF-kB": "炎症/NF-κB",
    "Hypoxia / oxidative stress": "缺氧/氧化应激",
    "Disc matrix homeostasis": "椎间盘基质稳态",
}


def clean_bibtex(value: str) -> str:
    """Remove BibTeX grouping braces while retaining the metadata text."""
    return re.sub(r"[{}]", "", value).replace("--", "-").strip()


def read_bibtex(path: Path) -> list[dict[str, str]]:
    """Read the project bibliography's one-field-per-line BibTeX format."""
    entries: list[dict[str, str]] = []
    current: dict[str, str] | None = None
    entry_start = re.compile(r"^@(\w+)\{([^,]+),")
    field_line = re.compile(r"^\s*([A-Za-z_]+)\s*=\s*\{(.*)\},?\s*$")

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        stripped = raw_line.strip()
        start = entry_start.match(stripped)
        if start:
            current = {"entry_type": start.group(1), "key": start.group(2)}
            continue
        if current is None:
            continue
        if stripped == "}":
            entries.append(current)
            current = None
            continue
        field = field_line.match(raw_line)
        if field:
            current[field.group(1).lower()] = field.group(2)

    if current is not None:
        raise ValueError(f"Unclosed BibTeX entry in {path}")
    if not entries:
        raise ValueError(f"No BibTeX entries found in {path}")
    return entries


def author_list(value: str) -> str:
    """Render BibTeX author names in a compact Vancouver-like form."""
    cleaned = clean_bibtex(value)
    if cleaned.startswith("NCBI Gene Expression Omnibus"):
        return "NCBI Gene Expression Omnibus"

    rendered: list[str] = []
    for author in cleaned.split(" and "):
        bits = [item.strip() for item in author.split(",", 1)]
        if len(bits) == 2:
            surname, given = bits
            initials = "".join(token[0] for token in re.findall(r"[A-Za-z]+", given))
            rendered.append(f"{surname} {initials}".strip())
        else:
            rendered.append(author)
    return ", ".join(rendered)


def vancouver_reference(number: int, entry: dict[str, str]) -> str:
    """Render the validated bibliography in a readable numeric reference style."""
    authors = author_list(entry.get("author", ""))
    title = clean_bibtex(entry.get("title", ""))
    year = clean_bibtex(entry.get("year", ""))
    doi = clean_bibtex(entry.get("doi", ""))
    url = clean_bibtex(entry.get("url", ""))
    if entry["entry_type"].lower() == "article":
        journal = clean_bibtex(entry.get("journal", ""))
        volume = clean_bibtex(entry.get("volume", ""))
        number_field = clean_bibtex(entry.get("number", ""))
        pages = clean_bibtex(entry.get("pages", ""))
        issue = f"({number_field})" if number_field else ""
        pagination = f":{pages}" if pages else ""
        doi_text = f" doi:{doi}." if doi else ""
        return f"{number}. {authors}. {title}. {journal}. {year};{volume}{issue}{pagination}.{doi_text}"

    how = clean_bibtex(entry.get("howpublished", ""))
    note = clean_bibtex(entry.get("note", ""))
    citation = f"{number}. {authors}. {title} [Internet]. {how}. {year}."
    if note:
        citation += f" {note}."
    if url:
        citation += f" Available from: {url}."
    return citation


def build_reference_block(entries: list[dict[str, str]]) -> str:
    return "\n\n".join(vancouver_reference(index, entry) for index, entry in enumerate(entries, start=1))


def extract_sections(markdown: str) -> dict[str, str]:
    """Return level-two Markdown sections without changing their nested headings."""
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in markdown.splitlines():
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
        elif current is not None:
            sections[current].append(line)
    return {name: "\n".join(lines).strip() for name, lines in sections.items()}


def replace_exact(text: str, old: str, new: str) -> str:
    if old not in text:
        raise ValueError(f"Expected text was not found: {old[:60]!r}")
    return text.replace(old, new, 1)


def enrich_english_sections(sections: dict[str, str]) -> dict[str, str]:
    """Apply formal-manuscript structure without changing audited results."""
    changed = {name: value.replace("NF-kB", "NF-κB") for name, value in sections.items()}
    neutral_terms = {
        "score-level directional analysis": "score-level directional description",
        "External NP support cohorts": "External NP cohorts",
        "directional-support": "external score-level",
        "count-level-support": "external dense-count",
        "normalized-count score-level\ndirection contrast": "normalized-count score-level\ndescriptive contrast",
    }
    for old, new in neutral_terms.items():
        changed = {name: value.replace(old, new) for name, value in changed.items()}
    methods = changed["Methods"]
    methods = replace_exact(
        methods,
        "### Cohorts and inference boundary",
        "### Study design, public datasets, and inference boundary",
    )
    methods = replace_exact(
        methods,
        "### Locked module scoring and effect displays",
        "### Locked program definition, scoring, and descriptive effect displays",
    )
    methods = replace_exact(
        methods,
        "### Separate exploratory standardized syntheses",
        "### Exploratory supplementary standardized syntheses",
    )
    methods = methods.rstrip() + "\n\n" + """

### Computational reproducibility

The default summary used 95% Welch intervals and 10,000 independent
donor or library bootstrap draws per contrast. Cohort-, compartment-, and
module-specific deterministic seeds were derived from root seed 20260814.
Scores with a mapped-gene fraction below 80% were excluded with an audit row.
The S7-S9 syntheses used R 4.4.1, `metafor` 4.8.0, and REML control
`maxiter = 10000`. Input, output, environment, and generator hashes are
indexed in Supplementary Table S6 and the submission-support manifest.
""".strip()

    methods = methods.replace(
        "Table 1 and Supplementary Table S2 describe cohort roles, recorded group\nstructures, observation keys, and identity checks.",
        "Table 1 and Supplementary Table S2 describe cohort roles, recorded group\nstructures, observation keys, and identity checks. Annulus fibrosus (AF) denotes\nthe outer annular compartment; extracellular matrix (ECM) denotes the matrix\nprogram label; and nuclear factor-kappa B (NF-κB) denotes the inflammatory\ntranscriptional response label.",
        1,
    )
    methods = methods.replace("FPKM representation", "fragments per kilobase of transcript per million mapped reads (FPKM) representation", 1)
    methods = methods.replace("TPM matrix", "transcripts per million (TPM) matrix", 1)
    methods = methods.replace("RPKM subset", "reads per kilobase per million mapped reads (RPKM) subset", 1)
    methods = methods.replace(
        "each module score was the mean mapped-gene `log1p(CPM)` value.",
        "counts per million (CPM) was the per-library normalization scale, and each module score was the mean mapped-gene `log1p(CPM)` value.",
        1,
    )
    methods = methods.replace(
        "each module score\nwas the mean mapped-gene `log1p(CPM)` value.",
        "counts per million (CPM) was the per-library normalization scale, and each module score\nwas the mean mapped-gene `log1p(CPM)` value.",
        1,
    )
    methods = methods.replace(
        "For every cohort, compartment, and module, we calculated the unweighted mean\ndifference between the recorded higher- and lower-severity groups.",
        "For every cohort, compartment, and module, we calculated the unweighted mean\ndifference between groups recorded as higher versus lower severity; positive values\ntherefore mean higher minus lower.",
        1,
    )
    methods = methods.replace("Leave-one-key-out analyses assessed", "Leave-one-key-out (LOKO) analyses assessed", 1)
    methods = methods.replace(
        "The primary\neffect was the heteroscedastic standardized mean difference (SMDH), fitted with\na REML random-effects model",
        "The primary\neffect was the heteroscedastic standardized mean difference (SMDH), fitted with\na restricted maximum likelihood (REML) random-effects model",
        1,
    )
    methods = methods.replace(
        "Four module-level HKSJ\np-values",
        "Four module-level Hartung-Knapp-Sidik-Jonkman (HKSJ) p-values",
        1,
    )
    changed["Methods"] = methods

    results = changed["Results"]
    results = replace_exact(
        results,
        "The default descriptive summary contains",
        "### Cohort disposition and reproducibility\n\nThe default descriptive summary contains",
    )
    results = replace_exact(
        results,
        "In NP, the hypoxia/oxidative-stress module",
        "### Default NP directional patterns\n\nIn NP, the hypoxia/oxidative-stress module",
    )
    results = replace_exact(
        results,
        "The separately audited GSE251686 sensitivity analysis",
        "### Isolated GSE251686 analysis\n\nThe separately audited GSE251686 sensitivity analysis",
    )
    results = replace_exact(
        results,
        "The separate S7 synthesis yielded",
        "### Exploratory supplementary syntheses\n\nThe separate S7 synthesis yielded",
    )
    results = replace_exact(
        results,
        "AF results were available only",
        "### Exploratory AF context and retained-cell threshold sensitivity\n\nAF results were available only",
    )
    changed["Results"] = results
    return changed


def document_front_matter(language: str, english_word_count: int) -> str:
    if language == "en":
        return f"""# {EN_TITLE}

## Title Page

**Article type:** Original Research

**Short title:** {EN_SHORT_TITLE}

**Authors:** [Author names to be inserted after author approval]

**Affiliations:** [Affiliations to be inserted after author approval]

**Corresponding author:** [Name, postal address, email, and telephone number to be inserted]

**Word count:** approximately {english_word_count:,} words, excluding title page, references, figure legends, and tables

**Figures and tables:** graphical abstract; 2 main figures; 2 main tables; 5 supplementary figures; 9 supplementary table packages
"""
    return f"""# {ZH_TITLE}

## 标题页

**文章类型：** 原创研究

**短题名：** {ZH_SHORT_TITLE}

**作者：** [待全体作者确认后补入]

**作者单位：** [待全体作者确认后补入]

**通讯作者：** [待补入姓名、通信地址、电子邮箱和电话]

**英文题名：** {EN_TITLE}

**图表：** 图形摘要；2 张主图；2 张主表；5 张补充图；9 组补充表
"""


def compose_english_markdown(project_root: Path, entries: list[dict[str, str]]) -> str:
    original = (project_root / "manuscript" / "04_manuscript_draft.md").read_text(encoding="utf-8")
    sections = enrich_english_sections(extract_sections(original))
    body_sections = [sections[name] for name in ("Abstract", "Introduction", "Methods", "Results", "Discussion", "Conclusions")]
    body_word_count = len(re.findall(r"\b[\w'-]+\b", " ".join(body_sections)))
    references = build_reference_block(entries)
    legends = sections["Table and Figure Legends"]
    supplementary_start = legends.find("**Supplementary Table S1.")
    if supplementary_start >= 0:
        legends = legends[supplementary_start:]
    legends = legends.replace(
        "The tables give S7 study-level effects, primary SMDH/REML/\nKnapp-Hartung results",
        "SMDH denotes the heteroscedastic standardized mean difference, and REML\ndenotes restricted maximum likelihood. The tables give S7 study-level effects,\nprimary SMDH/REML/Knapp-Hartung results",
        1,
    )
    legends = legends.replace(
        "The tables\ngive the separately packaged S8 study-level effects",
        "SMDH denotes the heteroscedastic standardized mean difference, and REML\ndenotes restricted maximum likelihood. The tables give the separately packaged S8\nstudy-level effects",
        1,
    )
    legends = legends.replace(
        "The tables give the S9 results",
        "SMDH denotes the heteroscedastic standardized mean difference, and REML\ndenotes restricted maximum likelihood. The tables give the S9 results",
        1,
    )
    legends = legends.replace(
        "The forest display gives S7 SMDH estimates and Knapp-Hartung\nintervals.",
        "The forest display gives S7 SMDH estimates and Knapp-Hartung intervals;\nSMDH is a heteroscedastic standardized mean difference and the display is\nnon-confirmatory.",
        1,
    )
    legends = legends.replace(
        "The forest\ndisplay gives S8 results",
        "The forest display gives S8 results",
        1,
    )
    legends = legends.replace(
        "The plotted\nSMDH estimates, intervals, and transparent p-values are non-confirmatory.",
        "The plotted SMDH estimates and intervals are non-confirmatory; SMDH is a\nheteroscedastic standardized mean difference, and HKSJ/BH p-values are reported\nfor transparency only.",
        1,
    )
    legends = legends.replace(
        "It demonstrates source-family sensitivity and does not add\nan independent validation cohort.",
        "It demonstrates source-family sensitivity and does not add an independent\nvalidation cohort. Axis ranges may differ among panels and effect magnitudes are\nnot directly comparable across processing scales.",
        1,
    )
    declarations = "\n\n".join(
        [
            "## Data Availability\n\n" + sections["Data Availability"],
            "## Code Availability\n\n" + sections["Code Availability"],
            "## Ethics Statement\n\n" + sections["Ethics Statement"],
            "## Informed Consent Statement\n\nNo new participant consent was obtained for this secondary analysis of de-identified publicly available data. The corresponding author must confirm whether any dataset-specific consent wording is required before submission.",
            "## Funding\n\n" + sections["Funding"],
            "## Competing Interests\n\n" + sections["Competing Interests"],
            "## Author Contributions\n\n" + sections["Author Contributions"],
            "## Acknowledgements\n\n[To be completed after author approval. Do not list contributors or support not confirmed by the authors.]",
        ]
    )
    table_and_figure_pages = """
## Main Tables

**Table 1. Cohort roles and inference boundaries in the default descriptive summary.**

[[TABLE:1]]

**Table 2. Cohort-specific NP module effects in the default descriptive summary.** The default layer contains 20 effects overall (16 NP effects and 4 exploratory AF effects); this table displays the 16 NP effects.

[[TABLE:2]]

## Main Figures

**Figure 1. Cohort-specific NP module-score differences and descriptive Welch 95% intervals.** Colors distinguish cohorts only; they do not encode effect magnitude.

[[FIGURE:figure_1]]

**Figure 2. NP cohort-specific directions and descriptive sign alignment.** Blue and orange encode positive and negative directions only; color does not encode effect magnitude.

[[FIGURE:figure_2]]

## Supplementary Material

Supplementary Tables S1-S4, S5a-S5b, S6, S7a-S7d, S8a-S8d, and S9a-S9d, together with Supplementary Figures S1-S5, are provided as separate submission files. Each S7-S9 artifact remains a separately packaged, exploratory, non-confirmatory synthesis. The graphical abstract should be uploaded separately if requested by the target journal.
""".strip()
    return "\n\n".join(
        [
            document_front_matter("en", body_word_count).strip(),
            "## Abstract\n\n" + EN_ABSTRACT,
            "## Keywords\n\n" + EN_KEYWORDS,
            "## Graphical Abstract\n\nPublic human IVDD cohorts were evaluated with donor or explicitly labelled presumed sample or library keys as the observation unit. Cells remained nested observations. Program definitions were locked before external scoring.\n\n[[FIGURE:graphical_abstract]]",
            "## Introduction\n\n" + sections["Introduction"],
            "## Methods\n\n" + sections["Methods"],
            "## Results\n\n" + sections["Results"],
            "## Discussion\n\n" + sections["Discussion"],
            "## Conclusions\n\n" + sections["Conclusions"],
            declarations,
            "## References\n\n" + references,
            "## Supplementary Table and Figure Legends\n\n" + legends,
            table_and_figure_pages,
        ]
    ) + "\n"


def compose_chinese_markdown(project_root: Path, entries: list[dict[str, str]]) -> str:
    template = (
        project_root / "manuscript" / "formal_submission" / "chinese_manuscript_body.md"
    ).read_text(encoding="utf-8")
    template = (
        template.replace("探索性 AF/NP 发现项目", "探索性 AF/NP 来源项目")
        .replace("外部 NP 支持队列", "外部 NP 来源队列")
        .replace("发现项目", "来源项目")
        .replace(
            "图示 SMDH、区间和透明报告的 p 值均为非确认性。",
            "图示 SMDH 和区间均为非确认性；HKSJ 与 BH p 值仅在补充表 S8b 中报告。",
        )
        .replace(
            "补充表 S1--S9 和补充图 S1--S5",
            "补充表 S1--S4、S5a--S5b、S6、S7a--S7d、S8a--S8d、S9a--S9d 以及补充图 S1--S5",
        )
    )
    sections = extract_sections(template)
    methods = sections["材料与方法"]
    methods = methods.replace(
        "表 1 和补充表 S2 描述队列角色、记录的分组结构、观察键和身份核对。",
        "表 1 和补充表 S2 描述队列角色、记录的分组结构、观察键和身份核对。本文中，\nAF 指纤维环（annulus fibrosus），ECM 指细胞外基质（extracellular matrix），\nNF-κB 指核因子 κB（nuclear factor-kappa B）。",
        1,
    )
    methods = methods.replace("FPKM 表示", "每百万比对读段每千碱基转录本片段数（FPKM）表示", 1)
    methods = methods.replace("TPM 矩阵", "每百万转录本（TPM）矩阵", 1)
    methods = methods.replace("RPKM 子集", "每百万比对读段每千碱基读段数（RPKM）子集", 1)
    methods = methods.replace("`log1p(CPM)` 值", "每百万计数（CPM）的 `log1p(CPM)` 值", 1)
    methods = methods.replace("逐一剔除观察键分析用于", "逐一剔除观察键（leave-one-key-out，LOKO）分析用于", 1)
    methods = methods.replace(
        "四个模块的 HKSJ p 值及其 Benjamini-Hochberg（BH）调整仅用于透明报告",
        "四个模块的 Hartung--Knapp--Sidik--Jonkman（HKSJ）p 值及其 Benjamini--Hochberg（BH）调整仅用于透明报告",
        1,
    )
    methods = methods.rstrip() + "\n\n" + """

### 计算可重复性

默认汇总使用 95% Welch 区间，并对每个对比进行 10,000 次独立的供体或文库键自助法重抽样。队列、部位和模块特异的确定性随机种子均由根种子 20260814 推导。映射基因比例低于 80% 的评分会连同审计记录一起排除。S7--S9 汇总使用 R 4.4.1、`metafor` 4.8.0 和 REML 控制参数 `maxiter = 10000`。输入、输出、环境和生成器哈希已在补充表 S6 与提交支持清单中索引。
""".strip()
    references = build_reference_block(entries)
    legends = sections["图表题注"]
    supplementary_start = legends.find("**补充表 S1.")
    if supplementary_start >= 0:
        legends = legends[supplementary_start:]
    legends = legends.replace(
        "这些表给出 S7 的研究层效应、主要 SMDH/REML/Knapp-Hartung 结果",
        "SMDH 指异方差标准化均数差，REML 指限制性最大似然法。这些表给出 S7 的研究层效应、主要 SMDH/REML/Knapp-Hartung 结果",
        1,
    )
    legends = legends.replace(
        "这些表在加入 GSE186542 和 GSE167931 FPKM 后给出",
        "SMDH 指异方差标准化均数差，REML 指限制性最大似然法。这些表在加入 GSE186542 和 GSE167931 FPKM 后给出",
        1,
    )
    legends = legends.replace(
        "这些表给出以 GSE245147 的原生对比替代 GSE167931 后的 S9 结果。",
        "SMDH 指异方差标准化均数差，REML 指限制性最大似然法。这些表给出以 GSE245147 的原生对比替代 GSE167931 后的 S9 结果。",
        1,
    )
    legends = legends.replace(
        "森林图给出 S7 SMDH 估计和 Knapp-Hartung 区间。",
        "森林图给出 S7 SMDH 估计和 Knapp-Hartung 区间；SMDH 指异方差标准化均数差。",
        1,
    )
    legends = legends.replace(
        "图示 SMDH、区间和透明报告的 p 值均为非确认性。",
        "图示 SMDH 和区间均为非确认性；SMDH 指异方差标准化均数差，HKSJ 与 BH p 值仅作透明报告。",
        1,
    )
    legends = legends.replace(
        "该图显示数据来源家族敏感性，不增加独立验证队列。",
        "该图显示数据来源家族敏感性，不增加独立验证队列。不同面板的横轴范围可能不同，跨处理尺度的效应量不应直接比较。",
        1,
    )
    table_and_figure_pages = """
## 主表

**表 1. 默认描述性汇总中各队列的角色与推断边界。**

[[TABLE:1]]

**表 2. 默认描述性汇总中各队列特异性的 NP 模块效应。** 默认层共包含 20 个效应（16 个 NP 效应和 4 个探索性 AF 效应）；本表展示其中 16 个 NP 效应。

[[TABLE:2]]

## 主图

**图 1. 各队列的 NP 模块分数差及描述性 Welch 95% 区间。** 颜色仅用于区分队列，不表示效应大小。

[[FIGURE:figure_1]]

**图 2. NP 队列特异性方向和描述性方向对齐。** 蓝色和橙色仅表示正向和负向，不表示效应大小。

[[FIGURE:figure_2]]
""".strip()
    declarations = "\n\n".join(
        [
            "## 数据可用性\n\n" + sections["数据可用性"],
            "## 代码可用性\n\n" + sections["代码可用性"],
            "## 伦理声明\n\n" + sections["伦理声明"],
            "## 知情同意\n\n本研究为已去标识化公共数据的二次分析，未新获取受试者知情同意。投稿前，通讯作者须确认是否需要披露任何数据集特异的知情同意信息。",
            "## 经费\n\n" + sections["经费"],
            "## 利益冲突\n\n" + sections["利益冲突"],
            "## 作者贡献\n\n" + sections["作者贡献"],
            "## 致谢\n\n[待全体作者确认后补充。不得列入未经确认的贡献者或支持信息。]",
        ]
    )
    return "\n\n".join(
        [
            document_front_matter("zh", 0).strip(),
            "## 摘要\n\n" + sections["摘要"],
            "## 关键词\n\n" + sections["关键词"],
            "## 图形摘要\n\n" + sections["图形摘要"] + "\n\n[[FIGURE:graphical_abstract]]",
            "## 引言\n\n" + sections["引言"],
            "## 材料与方法\n\n" + methods,
            "## 结果\n\n" + sections["结果"],
            "## 讨论\n\n" + sections["讨论"],
            "## 结论\n\n" + sections["结论"],
            declarations,
            "## 参考文献\n\n" + references,
            "## 补充表与图题注\n\n" + legends,
            table_and_figure_pages,
            "## 补充材料说明\n\n" + sections["补充材料说明"],
        ]
    ) + "\n"


def write_terminology_ledger(output_dir: Path) -> None:
    ledger = """# Terminology Ledger

| Canonical English term | Canonical Chinese term | Rule applied in both manuscripts |
|---|---|---|
| intervertebral disc degeneration (IVDD) | 椎间盘退变（IVDD） | Define once, then use IVDD. |
| nucleus pulposus (NP) | 髓核（NP） | Use NP only after first definition. |
| annulus fibrosus (AF) | 纤维环（AF） | AF results remain exploratory. |
| project-locked transcriptional program | 项目内预先锁定的转录模块 | Project-level lock, not prospective registration. |
| presumed donor/sample/library key | 推定供体/样本/文库键 | Never shorten to patient. |
| higher-recorded-severity minus lower-recorded-severity difference | 记录严重度较高组减较低组的差值 | Positive values indicate direction only. |
| descriptive Welch 95% interval | 描述性 Welch 95% 区间 | Do not call a zero-crossing interval non-significant. |
| leave-one-key-out analysis | 逐一剔除观察键分析 | A stability check, not cross-validation. |
| heteroscedastic standardized mean difference (SMDH) | 异方差标准化均数差（SMDH） | Used only in S7-S9. |
| non-confirmatory | 非确认性 | Never recast as validated or replicated. |
| source-family replacement sensitivity | 数据来源家族替换敏感性分析 | Does not add an independent cohort. |
"""
    (output_dir / "TERMINOLOGY_LEDGER.md").write_text(ledger, encoding="utf-8")


def write_delivery_readme(output_dir: Path) -> None:
    readme = """# Formal Manuscript Deliverables

These files are journal-neutral biomedical manuscript deliverables built from
the frozen project results. They use Vancouver-numbered references, A4 pages,
12-point body text, 1.5 line spacing, title pages, declarations, main tables,
and main figures.

The manuscripts are not yet journal-specific submission files. Author names,
affiliations, corresponding-author details, funding, conflicts of interest,
ethics determination, public-code repository URL, and archival DOI remain
explicit placeholders pending author confirmation. The S7-S9 analyses are
retained as non-confirmatory supplementary analyses and do not alter the
default 20-effect descriptive analysis.
"""
    (output_dir / "README.md").write_text(readme, encoding="utf-8")


def set_run_font(run, latin_font: str, cjk_font: str, size: Pt | None = None, bold: bool | None = None) -> None:
    run.font.name = latin_font
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), cjk_font)


def set_style_font(style, latin_font: str, cjk_font: str, size: Pt, bold: bool = False) -> None:
    style.font.name = latin_font
    style.font.size = size
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), cjk_font)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, latin_font: str, cjk_font: str, size: Pt, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, latin_font, cjk_font, size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph, latin_font: str, cjk_font: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, latin_font, cjk_font, Pt(9))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_section(section, landscape: bool, latin_font: str, cjk_font: str, short_title: str) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54 if not landscape else 1.5)
    section.right_margin = Cm(2.54 if not landscape else 1.5)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Cm(29.7), Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(short_title)
    set_run_font(run, latin_font, cjk_font, Pt(8))
    run.font.color.rgb = RGBColor(89, 89, 89)
    add_page_number(section.footer.paragraphs[0], latin_font, cjk_font)


def add_inline_runs(paragraph, text: str, latin_font: str, cjk_font: str, size: Pt, default_bold: bool = False) -> None:
    """Render a small Markdown subset used by the manuscript source."""
    token = re.compile(r"(\*\*.*?\*\*|`.*?`|\*[^*]+?\*)")
    position = 0
    for match in token.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, latin_font, cjk_font, size, default_bold)
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            set_run_font(run, latin_font, cjk_font, size, True)
        elif value.startswith("`"):
            run = paragraph.add_run(value[1:-1])
            set_run_font(run, "Consolas", cjk_font, size, default_bold)
        else:
            run = paragraph.add_run(value[1:-1])
            set_run_font(run, latin_font, cjk_font, size, default_bold)
            run.italic = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, latin_font, cjk_font, size, default_bold)


def add_body_paragraph(doc: Document, text: str, latin_font: str, cjk_font: str, reference: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    if reference:
        paragraph.paragraph_format.left_indent = Cm(0.63)
        paragraph.paragraph_format.first_line_indent = Cm(-0.63)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(3)
    add_inline_runs(paragraph, text, latin_font, cjk_font, Pt(12 if not reference else 9))


def add_caption(doc: Document, text: str, latin_font: str, cjk_font: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    add_inline_runs(paragraph, text, latin_font, cjk_font, Pt(9))


def add_heading(doc: Document, title: str, level: int, latin_font: str, cjk_font: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title)
    size = Pt(14 if level == 1 else 12)
    set_run_font(run, latin_font, cjk_font, size, True)


def concise_table_rows(project_root: Path, language: str) -> tuple[list[str], list[list[str]], list[str], list[list[str]]]:
    with (project_root / "results" / "tables" / "table_1_current_cohort_roles.csv").open(encoding="utf-8", newline="") as handle:
        table_1 = list(csv.DictReader(handle))
    with (project_root / "results" / "tables" / "table_2_np_module_effects.csv").open(encoding="utf-8", newline="") as handle:
        table_2 = list(csv.DictReader(handle))

    if language == "en":
        cohort_names = {
            "GSE230809 AF + NP\nexploratory (AF 3 vs 10; NP 3 vs 8)": "GSE230809 AF/NP parent project, exploratory (AF 3 vs 10; NP 3 vs 8)",
            "GSE244889 NP\ndirectional support (4 vs 3)": "GSE244889 external NP score-level cohort (4 vs 3)",
            "GSE153066 NP\ncount-level support (8 vs 8)": "GSE153066 external NP dense-count cohort (8 vs 8)",
            "GSE165722 NP\nscore-level support (4 vs 4)": "GSE165722 external NP normalized-count cohort (4 vs 4)",
            "GSE230809 NP\nexploratory (3 vs 8)": "GSE230809 NP parent-project comparison, exploratory (3 vs 8)",
        }
        headers_1 = ["Cohort and role", "Recorded group structure", "Observation unit", "Inference boundary"]
        rows_1 = [
            [
                cohort_names.get(row["Cohort and role"], row["Cohort and role"].replace("\n", " ")),
                row["Recorded group structure"],
                row["Observation key"],
                row["Interpretation boundary"],
            ]
            for row in table_1
        ]
        headers_2 = ["Cohort", "Module", "n (lower vs higher)", "Difference", "Welch 95% CI", "Bootstrap 95% CI", "LOKO retention"]
        rows_2 = [
            [
                cohort_names.get(row["Cohort and role"], row["Cohort and role"].replace("\n", " ")),
                MODULE_LABELS_EN.get(row["Pre-specified module"], row["Pre-specified module"]),
                row["Lower vs higher group n"],
                row["Higher-minus-lower score difference"],
                row["Welch 95% interval"],
                row["Bootstrap 95% interval"],
                row["LODO direction retention (fraction)"],
            ]
            for row in table_2
        ]
        return headers_1, rows_1, headers_2, rows_2

    cohort_names = {
        "GSE230809 AF + NP\nexploratory (AF 3 vs 10; NP 3 vs 8)": "GSE230809 AF/NP 探索性父项目（AF 3 对 10；NP 3 对 8）",
        "GSE244889 NP\ndirectional support (4 vs 3)": "GSE244889 NP 分数层面队列（4 对 3）",
        "GSE153066 NP\ncount-level support (8 vs 8)": "GSE153066 NP 计数层面队列（8 对 8）",
        "GSE165722 NP\nscore-level support (4 vs 4)": "GSE165722 NP 标准化计数分数层面队列（4 对 4）",
        "GSE230809 NP\nexploratory (3 vs 8)": "GSE230809 NP 探索性父项目（3 对 8）",
    }
    key_names = {
        "donor_id": "供体标识",
        "presumed donor/library key": "推定供体/文库键",
        "presumed donor-level sample key": "推定供体层面样本键",
    }
    group_names = {
        "healthy n=3 per compartment; diseased AF n=10, NP n=8": "每个部位健康组 n=3；退变 AF n=10、NP n=8",
        "MDD n=4; SDD n=3": "MDD n=4；SDD n=3",
        "relatively normal n=8; degenerated n=8": "相对正常 n=8；退变 n=8",
        "mild n=4; severe n=4": "轻度 n=4；重度 n=4",
    }
    short_boundaries = {
        "donor_id": "单一父项目；年龄与疾病状态混杂；仅探索性展示。",
        "presumed donor/library key": "推定观察键；仅描述性分数方向。",
        "presumed donor-level sample key": "标准化计数；仅分数层面方向。",
    }
    headers_1 = ["队列及角色", "记录的分组结构", "观察单位", "推断边界"]
    rows_1 = [
        [
            cohort_names.get(row["Cohort and role"], row["Cohort and role"].replace("\n", " ")),
            group_names.get(row["Recorded group structure"], row["Recorded group structure"]),
            key_names.get(row["Observation key"], row["Observation key"]),
            short_boundaries.get(row["Observation key"], "所有当前队列均不具确认性资格。"),
        ]
        for row in table_1
    ]
    headers_2 = ["队列", "模块", "n（低对高）", "差值", "Welch 95% 区间", "自助法 95% 区间", "逐一剔除观察键方向保留"]
    rows_2 = [
        [
            cohort_names.get(row["Cohort and role"], row["Cohort and role"].replace("\n", " ")),
            MODULE_LABELS_ZH.get(row["Pre-specified module"], row["Pre-specified module"]),
            row["Lower vs higher group n"],
            row["Higher-minus-lower score difference"],
            row["Welch 95% interval"],
            row["Bootstrap 95% interval"],
            row["LODO direction retention (fraction)"],
        ]
        for row in table_2
    ]
    return headers_1, rows_1, headers_2, rows_2


def add_table(doc: Document, headers: list[str], rows: list[list[str]], latin_font: str, cjk_font: str) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "1F4E78")
        set_cell_text(cell, header, latin_font, cjk_font, Pt(8), True)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if len(table.rows) % 2 == 0:
                shade_cell(cells[index], "EAF2F8")
            set_cell_text(cells[index], value, latin_font, cjk_font, Pt(7.5))
    widths = (
        [Cm(4.7), Cm(5.0), Cm(4.2), Cm(10.6)]
        if len(headers) == 4
        else [Cm(3.7), Cm(3.6), Cm(2.2), Cm(2.1), Cm(4.1), Cm(4.1), Cm(4.7)]
    )
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_properties.append(header_marker)
    doc.add_paragraph()


def add_figure(
    doc: Document,
    project_root: Path,
    figure_id: str,
    latin_font: str,
    cjk_font: str,
    width_inches: float = 6.2,
) -> None:
    figure_paths = {
        "graphical_abstract": project_root / "results" / "graphical_abstract" / "graphical_abstract_cohort_aware_ivdd.png",
        # Formal manuscript figures are kept separate from the historical
        # deliverables so their neutral cohort labels cannot alter audit files.
        "figure_1": project_root / "manuscript" / "formal_submission" / "formal_figures" / "figure_1_np_cohort_module_effects.png",
        "figure_2": project_root / "manuscript" / "formal_submission" / "formal_figures" / "figure_2_np_direction_alignment.png",
    }
    path = figure_paths[figure_id]
    if not path.exists() and figure_id in {"figure_1", "figure_2"}:
        path = project_root / "results" / "figures" / path.name
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    doc.add_paragraph()


def render_docx(markdown: str, destination: Path, project_root: Path, language: str) -> None:
    latin_font, cjk_font = ("Times New Roman", "SimSun") if language == "en" else ("Times New Roman", "SimSun")
    short_title = EN_SHORT_TITLE if language == "en" else ZH_SHORT_TITLE
    doc = Document()
    configure_section(doc.sections[0], False, latin_font, cjk_font, short_title)
    doc.sections[0].different_first_page_header_footer = True
    normal = doc.styles["Normal"]
    set_style_font(normal, latin_font, cjk_font, Pt(12))
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    if "Caption" not in doc.styles:
        doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)

    tables = concise_table_rows(project_root, language)
    paragraph_lines: list[str] = []
    in_references = False
    landscape_tables = False
    main_figures_started = False
    main_figure_count = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = " ".join(item.strip() for item in paragraph_lines).strip()
            if text:
                if text.startswith("**"):
                    add_caption(doc, text, latin_font, cjk_font)
                else:
                    add_body_paragraph(doc, text, latin_font, cjk_font, in_references)
            paragraph_lines = []

    def is_figure_caption() -> bool:
        if not paragraph_lines:
            return False
        text = " ".join(item.strip() for item in paragraph_lines).strip()
        return text.startswith("**Figure ") or text.startswith("**图 ")

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("# "):
            flush_paragraph()
            title = line[2:].strip()
            doc.add_paragraph()
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(72)
            paragraph.paragraph_format.space_after = Pt(18)
            add_inline_runs(paragraph, title, latin_font, cjk_font, Pt(18), True)
            continue
        if line.startswith("## "):
            flush_paragraph()
            heading = line[3:].strip()
            if heading == "Abstract" or heading == "\u6458\u8981":
                doc.add_page_break()
            if heading == "Graphical Abstract" or heading == "\u56fe\u5f62\u6458\u8981":
                doc.add_page_break()
            if heading in {"Main Tables", "主表"}:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, True, latin_font, cjk_font, short_title)
                landscape_tables = True
            elif heading in {"Main Figures", "主图"}:
                main_figures_started = True
                if landscape_tables:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                    configure_section(section, True, latin_font, cjk_font, short_title)
                else:
                    doc.add_page_break()
            elif heading in {"Supplementary Material", "补充材料说明"} and landscape_tables:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, False, latin_font, cjk_font, short_title)
                landscape_tables = False
            elif heading in {"References", "参考文献"}:
                doc.add_page_break()
                in_references = True
            else:
                in_references = False
            add_heading(doc, heading, 1, latin_font, cjk_font)
            continue
        if line.startswith("### "):
            flush_paragraph()
            add_heading(doc, line[4:].strip(), 2, latin_font, cjk_font)
            continue
        if line.startswith("[[TABLE:") and line.endswith("]]" ):
            flush_paragraph()
            number = line[8:-2]
            if number == "1":
                add_table(doc, tables[0], tables[1], latin_font, cjk_font)
            elif number == "2":
                add_table(doc, tables[2], tables[3], latin_font, cjk_font)
            else:
                raise ValueError(f"Unknown table marker: {line}")
            continue
        if line.startswith("[[FIGURE:") and line.endswith("]]" ):
            figure_id = line[9:-2]
            is_main_figure = figure_id in {"figure_1", "figure_2"}
            first_main_figure_on_section = is_main_figure and main_figures_started and main_figure_count == 0
            # Main figures already flow within a landscape section. Forcing a
            # second break after Figure 1 creates an empty intervening page in
            # Word/WPS; the caption's keep-with-next setting handles the wrap.
            needs_page_break = (
                figure_id != "graphical_abstract"
                and not first_main_figure_on_section
                and not (is_main_figure and landscape_tables)
            )
            if needs_page_break:
                doc.add_page_break()
            flush_paragraph()
            # Figure 1 shares its first landscape page with the section title
            # and caption, whereas Figure 2 does not need that extra headroom.
            # This keeps each caption physically attached to its figure.
            if figure_id == "figure_1" and landscape_tables:
                width_inches = 8.5
            else:
                width_inches = 9.2 if is_main_figure and landscape_tables else 6.2
            add_figure(doc, project_root, figure_id, latin_font, cjk_font, width_inches)
            if is_main_figure:
                main_figure_count += 1
            continue
        if not line.strip():
            # Keep a main-figure caption attached to the following figure
            # marker; otherwise the page break would separate the caption
            # from its image.
            if not is_figure_caption():
                flush_paragraph()
            continue
        if in_references and re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            add_body_paragraph(doc, line, latin_font, cjk_font, reference=True)
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    doc.save(destination)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-root", type=Path, default=Path("."))
    parser.add_argument("--output-dir", type=Path, default=Path("manuscript/formal_submission"))
    args = parser.parse_args()

    project_root = args.project_root.resolve()
    output_dir = (project_root / args.output_dir).resolve() if not args.output_dir.is_absolute() else args.output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    entries = read_bibtex(project_root / "manuscript" / "references.bib")
    if len(entries) != 27:
        raise ValueError(f"Expected 27 validated references, found {len(entries)}")

    english_markdown = compose_english_markdown(project_root, entries)
    chinese_markdown = compose_chinese_markdown(project_root, entries)
    en_markdown_path = output_dir / "IVDD_cohort_aware_manuscript_EN.md"
    zh_markdown_path = output_dir / "IVDD_cohort_aware_manuscript_ZH.md"
    en_markdown_path.write_text(english_markdown, encoding="utf-8")
    zh_markdown_path.write_text(chinese_markdown, encoding="utf-8")
    write_terminology_ledger(output_dir)
    write_delivery_readme(output_dir)

    render_docx(english_markdown, output_dir / "IVDD_cohort_aware_manuscript_EN.docx", project_root, "en")
    render_docx(chinese_markdown, output_dir / "IVDD_cohort_aware_manuscript_ZH.docx", project_root, "zh")
    print(f"Wrote formal manuscript sources and DOCX files to {output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
